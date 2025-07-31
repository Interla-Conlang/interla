"""
Generate dictionaries for Interla (from Interla to other languages, and reversed).
Creates printable DOCX dictionaries with 3-column layout.
"""

import os
import pickle
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from tqdm import tqdm
from tqdm.contrib.concurrent import process_map

from assign_spellings_common import get_data_from_wiktionary
from lan_freqs import FREQ_DICT
from logging_config import logger
from utils import ALL_VALID_LANGUAGES, INTERLA_TO_IPA, LANG_WEIGHTS


def interla_to_ipa(interla_word: str) -> str:
    """Convert Interla word to IPA notation using the alphabet mapping."""
    ipa_chars = []
    i = 0
    while i < len(interla_word):
        # Check for two-character combinations first (like 'ch', 'ng')
        if i < len(interla_word) - 1:
            two_char = interla_word[i : i + 2]
            if two_char in INTERLA_TO_IPA:
                ipa_chars.append(INTERLA_TO_IPA[two_char])
                i += 2
                continue

        # Check single character
        if interla_word[i] in INTERLA_TO_IPA:
            ipa_chars.append(INTERLA_TO_IPA[interla_word[i]])
        else:
            ipa_chars.append(interla_word[i])  # Keep unknown characters as-is
        i += 1

    return "".join(ipa_chars)


def create_three_column_section(doc, entries: List[Tuple[str, str, List[str]]]):
    """Create entries using Word's built-in 3-column layout for continuous text flow."""
    if not entries:
        return

    # Set up the current section for 3 columns
    section = doc.sections[-1]

    # Configure page margins for better column layout
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Set up 3 columns using the section properties
    sectPr = section._sectPr

    # Remove existing column settings
    cols_list = sectPr.xpath(".//w:cols")
    for cols in cols_list:
        sectPr.remove(cols)

    # Add new column settings
    from docx.oxml import parse_xml

    cols_xml = """<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:num="3" w:space="360"/>"""
    cols_element = parse_xml(cols_xml)
    sectPr.append(cols_element)

    # Add all entries as continuous paragraphs
    for word, ipa, translations in tqdm(entries, desc="Adding entries"):
        p = doc.add_paragraph()

        # Add word in bold
        run = p.add_run(word)  # FIXME: 13% of time
        run.bold = True  # FIXME: 23% of time

        # Add IPA if provided (on same line, in italics)
        if ipa:
            ipa_run = p.add_run(f" /{ipa}/")
            ipa_run.italic = True  # FIXME: 10% of time

        # Add translations on new line (join multiple translations with commas)
        translation_text = ", ".join(translations)
        p.add_run(f"\n{translation_text}")  # FIXME: 20% of time

        # Add paragraph spacing for better readability
        p.space_after = Inches(0.08)


def create_dictionary_docx(
    lang_to_interla: Dict[str, List[str]],
    interla_to_lang: Dict[str, List[str]],
    language_code: str,
) -> None:
    """Create a DOCX dictionary for a specific language."""
    doc = Document()

    # Title
    title = doc.add_heading(f"Interla-{language_code.upper()} Dictionary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Part 1: Original Language to Interla
    doc.add_heading(f"{language_code.upper()} to Interla", level=1)

    # Sort entries alphabetically by original language word
    lang_entries = []
    for orig_word, interla_words in sorted(lang_to_interla.items()):
        ipa = ""  # Original language words don't have IPA in this context
        lang_entries.append((orig_word, ipa, interla_words))

    if lang_entries:
        create_three_column_section(doc, lang_entries)
    else:
        doc.add_paragraph("No entries found.")

    # Add a new section with single column for the page break and heading
    doc.add_section()

    # Reset to single column for the heading
    section = doc.sections[-1]
    sectPr = section._sectPr
    cols_list = sectPr.xpath(".//w:cols")
    for cols in cols_list:
        sectPr.remove(cols)

    # Part 2: Interla to Original Language
    doc.add_heading(f"Interla to {language_code.upper()}", level=1)

    # Sort entries alphabetically by Interla word
    interla_entries = []
    for interla_word, orig_words in sorted(interla_to_lang.items()):
        ipa = interla_to_ipa(interla_word)
        interla_entries.append((interla_word, ipa, orig_words))

    if interla_entries:
        create_three_column_section(doc, interla_entries)
    else:
        doc.add_paragraph("No entries found.")

    # Save the document
    filename = f"output/dictionaries/dictionary_interla_{language_code}.docx"
    doc.save(filename)
    logger.info(f"Dictionary saved as {filename}")


def gen_single_dictionary(language_code: str) -> None:
    """Generate DOCX dictionary for a single language."""
    logger.info(f"Starting dictionary generation for {language_code}...")

    # Load data
    int_anon_tokens_coocurrences, all_y2normWord, all_y2word, _ = (
        get_data_from_wiktionary()
    )

    interla_vocab_path = "output/interla_vocab.pkl"

    if not os.path.exists(interla_vocab_path):
        logger.error(f"Interla vocabulary file not found at {interla_vocab_path}")
        return

    with open(interla_vocab_path, "rb") as f:
        vocab: Dict[str, int] = pickle.load(f)

    if language_code not in all_y2word:
        logger.error(f"Language {language_code} not found in available languages")
        logger.info(f"Available languages: {sorted(all_y2word.keys())}")
        return

    # Build dictionaries for this language
    lang_to_interla = {}  # Original language -> List[Interla]
    interla_to_lang = {}  # Interla -> List[Original language]

    y2word = all_y2word[language_code]

    freq_dict = FREQ_DICT[language_code] if language_code in FREQ_DICT else None
    for int_orth_token, int_anon_token in vocab.items():
        assoc_words = int_anon_tokens_coocurrences.get(int_anon_token, {})
        if language_code in assoc_words:
            y_id = assoc_words[language_code]
            word = y2word[y_id]

            if freq_dict is None or word in freq_dict:
                # Store both directions, handling multiple translations
                if word not in lang_to_interla:
                    lang_to_interla[word] = []
                if int_orth_token not in lang_to_interla[word]:
                    lang_to_interla[word].append(int_orth_token)

                if int_orth_token not in interla_to_lang:
                    interla_to_lang[int_orth_token] = []
                if word not in interla_to_lang[int_orth_token]:
                    interla_to_lang[int_orth_token].append(word)

    if lang_to_interla:
        logger.info(
            f"Creating dictionary for {language_code} with {len(lang_to_interla)} entries"
        )
        create_dictionary_docx(lang_to_interla, interla_to_lang, language_code)
    else:
        logger.warning(f"No entries found for language: {language_code}")


def create_dictionary_worker(
    args: Tuple[Dict[str, List[str]], Dict[str, List[str]], str],
) -> str:
    """Worker function to create a dictionary DOCX file for parallel processing."""
    lang_to_interla, interla_to_lang, dict_lang = args

    try:
        if lang_to_interla:
            create_dictionary_docx(lang_to_interla, interla_to_lang, dict_lang)
            return f"Successfully created dictionary for {dict_lang} with {len(lang_to_interla)} entries"
        else:
            return f"Warning: No entries found for language {dict_lang}"

    except Exception as e:
        return f"Error creating dictionary for {dict_lang}: {str(e)}"


def gen_dictionaries() -> None:
    """Generate DOCX dictionaries for all languages using parallel processing."""
    logger.info("Starting dictionary generation with parallel processing...")

    # Load data once in the main process
    int_anon_tokens_cooccurrences, _, all_y2word, _ = get_data_from_wiktionary()

    interla_vocab_path = "output/interla_vocab.pkl"
    if not os.path.exists(interla_vocab_path):
        logger.error(f"Interla vocabulary file not found at {interla_vocab_path}")
        return

    with open(interla_vocab_path, "rb") as f:
        vocab: Dict[str, int] = pickle.load(f)

    # Prepare dictionary data for all languages
    logger.info("Preparing dictionary data for all languages...")
    dictionary_args = []

    # Sort ALL_VALID_LANGUAGES by language weight (descending)
    sorted_langs = sorted(
        ALL_VALID_LANGUAGES,
        key=lambda lang: LANG_WEIGHTS.get(lang, 0),
        reverse=True,
    )
    for dict_lang in tqdm(sorted_langs, desc="Preparing data"):
        # Build dictionaries for this language
        lang_to_interla = {}  # Original language -> List[Interla]
        interla_to_lang = {}  # Interla -> List[Original language]

        if dict_lang not in all_y2word:
            logger.warning(f"Language {dict_lang} not found in available languages")
            continue
        y2word = all_y2word[dict_lang]
        freq_dict = FREQ_DICT[dict_lang] if dict_lang in FREQ_DICT else None
        for int_orth_token, int_anon_token in vocab.items():
            assoc_words = int_anon_tokens_cooccurrences.get(int_anon_token, {})
            if dict_lang in assoc_words:
                y_id = assoc_words[dict_lang]
                word = y2word[y_id]

                if (
                    freq_dict is None or word.lower() in freq_dict
                ):  # restrict to frequent words
                    # Store both directions, handling multiple translations
                    if word not in lang_to_interla:
                        lang_to_interla[word] = []
                    if int_orth_token not in lang_to_interla[word]:
                        lang_to_interla[word].append(int_orth_token)

                    if int_orth_token not in interla_to_lang:
                        interla_to_lang[int_orth_token] = []
                    if word not in interla_to_lang[int_orth_token]:
                        interla_to_lang[int_orth_token].append(word)

        # Add to arguments list for parallel processing
        dictionary_args.append((lang_to_interla, interla_to_lang, dict_lang))

    results = process_map(
        create_dictionary_worker,
        dictionary_args,
        max_workers=10,
        desc="Creating DOCX files",
    )

    # Log results
    for result in results:
        if result.startswith("Successfully"):
            logger.info(result)
        elif result.startswith("Warning"):
            logger.warning(result)
        else:
            logger.error(result)

    logger.info("Dictionary generation completed!")


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        # Generate dictionary for specific language
        lang_code = sys.argv[1]
        gen_single_dictionary(lang_code)
    else:
        # Generate dictionaries for all languages
        gen_dictionaries()
